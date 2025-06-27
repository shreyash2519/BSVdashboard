# Paste this entire code into your `psbsv.py`

import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from lifelines import KaplanMeierFitter, CoxPHFitter, statistics
import base64, time, tempfile
from docx import Document
from docx.shared import Inches

# === Splash Screen (3 seconds fullscreen) ===
def get_base64_image(image_path):
    with open(image_path, "rb") as f:
        return base64.b64encode(f.read()).decode()

if "show_main" not in st.session_state:
    st.session_state.show_main = False
if not st.session_state.show_main:
    img_base64 = get_base64_image("assets/logo.png")
    st.markdown(f"""
        <style>
        .splash {{
            position: fixed; top: 0; left: 0; right: 0; bottom: 0;
            background-color: black;
            display: flex; justify-content: center; align-items: center;
            z-index: 9999;
        }}
        .splash img {{
            width: 100vw; height: 100vh; object-fit: cover;
        }}
        </style>
        <div class="splash"><img src="data:image/png;base64,{img_base64}" /></div>
    """, unsafe_allow_html=True)
    time.sleep(3)
    st.session_state.show_main = True
    st.rerun()

# === App Main Title ===
st.title("📊 PS BioStatView:- Insights in Oncology")

# === Upload Dataset ===
with st.sidebar:
    st.markdown("### 📁 Upload your dataset")
    uploaded_file = st.file_uploader("Upload file", type=["csv", "xlsx", "xls", "json"])

if uploaded_file:
    file_name = uploaded_file.name.split(".")[0]
    ext = uploaded_file.name.split(".")[-1].lower()
    if ext == "csv":
        df = pd.read_csv(uploaded_file)
    elif ext in ["xlsx", "xls"]:
        df = pd.read_excel(uploaded_file)
    elif ext == "json":
        df = pd.read_json(uploaded_file)
    else:
        st.error("Unsupported file format.")
        st.stop()

    st.session_state.df = df
    st.subheader("📊 Uploaded Data Preview")
    st.dataframe(df.head())

    with st.sidebar:
        analysis = st.selectbox("📌 Select analysis", ["Descriptive Stats", "Kaplan-Meier + Log-Rank Test", "Cox Proportional Hazards"])
        chart_types = st.multiselect("📈 Select Charts", ["Bar Chart", "Pie Chart", "Histogram", "Box Plot", "Scatter Plot"])
        download_button = st.button("⬇️ Download Word Report")

    # === Descriptive Stats ===
    if analysis == "Descriptive Stats":
        st.session_state["ran_desc"] = True
        cat_cols = df.select_dtypes(include='object').columns
        num_cols = df.select_dtypes(include=np.number).columns

        if len(cat_cols) > 0:
            st.subheader("📋 Categorical Summary")
            cat_df = pd.DataFrame([{
                "Variable": col,
                "Unique": df[col].nunique(),
                "Missing": df[col].isnull().sum()
            } for col in cat_cols])
            st.session_state["cat_table"] = cat_df
            st.dataframe(cat_df)

        if len(num_cols) > 0:
            st.subheader("📈 Numerical Summary")
            num_summary = []
            for col in num_cols:
                desc = df[col].describe()
                iqr = desc['75%'] - desc['25%']
                outliers = df[(df[col] < desc['25%'] - 1.5 * iqr) | (df[col] > desc['75%'] + 1.5 * iqr)].shape[0]
                num_summary.append({
                    "Variable": col,
                    "Mean": round(desc['mean'], 2),
                    "Median": round(df[col].median(), 2),
                    "Std": round(desc['std'], 2),
                    "Min": round(desc['min'], 2),
                    "Max": round(desc['max'], 2),
                    "Outliers": outliers
                })
            num_df = pd.DataFrame(num_summary)
            st.session_state["num_table"] = num_df
            st.dataframe(num_df)

    # === Kaplan-Meier ===
    elif analysis == "Kaplan-Meier + Log-Rank Test":
        st.session_state["ran_km"] = True
        st.subheader("📈 Kaplan-Meier Curve + Log-Rank Test")
        duration = st.selectbox("Duration Column", df.columns)
        event = st.selectbox("Event Column (0=censored, 1=event)", df.columns)
        group = st.selectbox("Group Column (optional)", ["None"] + list(df.columns))

        kmf = KaplanMeierFitter()
        fig_km, ax = plt.subplots()

        if group == "None":
            kmf.fit(df[duration], event_observed=df[event])
            kmf.plot_survival_function(ax=ax)
        else:
            groups = df[group].dropna().unique()
            for g in groups:
                kmf.fit(df[df[group] == g][duration], df[df[group] == g][event], label=str(g))
                kmf.plot_survival_function(ax=ax)
            if len(groups) == 2:
                d1 = df[df[group] == groups[0]]
                d2 = df[df[group] == groups[1]]
                result = statistics.logrank_test(d1[duration], d2[duration], event_observed_A=d1[event], event_observed_B=d2[event])
                st.session_state["last_pvalue"] = result.p_value
                st.markdown(f"**Log-Rank p-value:** {result.p_value:.4f}")

        st.session_state["km_plot"] = fig_km
        st.pyplot(fig_km)

    # === Cox Model ===
    elif analysis == "Cox Proportional Hazards":
        st.session_state["ran_cox"] = True
        st.subheader("📌 Cox Proportional Hazards")
        duration = st.selectbox("Duration", df.columns)
        event = st.selectbox("Event", df.columns)
        covs = st.multiselect("Covariates", [c for c in df.columns if c not in [duration, event]])
        if covs:
            cph = CoxPHFitter()
            cph.fit(df[[duration, event] + covs].dropna(), duration_col=duration, event_col=event)
            st.session_state["cox_summary"] = cph.summary
            fig_cox = cph.plot()
            st.session_state["cox_plot"] = fig_cox.figure
            st.write(cph.summary)
            st.pyplot(fig_cox.figure)

    # === Chart Section ===
    chart_figures = []
    if chart_types:
        st.subheader("📊 Chart Results")
        for i, chart_type in enumerate(chart_types):
            fig, ax = plt.subplots()
            if chart_type == "Bar Chart":
                col = st.selectbox("Bar Chart Variable", df.select_dtypes(include='object').columns, key=f"bar_{i}")
                df[col].value_counts().plot(kind='bar', ax=ax)
                ax.set_title(f"Bar Chart of {col}")
            elif chart_type == "Pie Chart":
                col = st.selectbox("Pie Chart Variable", df.select_dtypes(include='object').columns, key=f"pie_{i}")
                df[col].value_counts().plot(kind='pie', autopct='%1.1f%%', ax=ax)
                ax.set_ylabel('')
                ax.set_title(f"Pie Chart of {col}")
            elif chart_type == "Histogram":
                col = st.selectbox("Histogram Variable", df.select_dtypes(include='number').columns, key=f"hist_{i}")
                sns.histplot(df[col].dropna(), kde=True, ax=ax)
                ax.set_title(f"Histogram of {col}")
            elif chart_type == "Box Plot":
                y = st.selectbox("Box Numeric", df.select_dtypes(include='number').columns, key=f"box_y_{i}")
                x = st.selectbox("Box Category", df.select_dtypes(include='object').columns, key=f"box_x_{i}")
                sns.boxplot(x=x, y=y, data=df, ax=ax)
                ax.set_title(f"Boxplot of {y} by {x}")
            elif chart_type == "Scatter Plot":
                x = st.selectbox("Scatter X", df.select_dtypes(include='number').columns, key=f"scatter_x_{i}")
                y = st.selectbox("Scatter Y", df.select_dtypes(include='number').columns, key=f"scatter_y_{i}")
                sns.scatterplot(x=df[x], y=df[y], ax=ax)
                ax.set_title(f"Scatter: {x} vs {y}")
            chart_figures.append((chart_type, fig))
            st.pyplot(fig)
        st.session_state["charts"] = chart_figures

    # === Generate Word Report ===
    if download_button:
        doc = Document()
        doc.add_heading("PS BioStatView Report", 0)

        # Descriptive
        if st.session_state.get("ran_desc"):
            if "cat_table" in st.session_state:
                doc.add_heading("Categorical Summary", level=1)
                df_cat = st.session_state["cat_table"]
                table = doc.add_table(rows=1, cols=len(df_cat.columns))
                for i, col in enumerate(df_cat.columns):
                    table.cell(0, i).text = col
                for row in df_cat.itertuples(index=False):
                    cells = table.add_row().cells
                    for i, val in enumerate(row):
                        cells[i].text = str(val)
            if "num_table" in st.session_state:
                doc.add_heading("Numerical Summary", level=1)
                df_num = st.session_state["num_table"]
                table = doc.add_table(rows=1, cols=len(df_num.columns))
                for i, col in enumerate(df_num.columns):
                    table.cell(0, i).text = col
                for row in df_num.itertuples(index=False):
                    cells = table.add_row().cells
                    for i, val in enumerate(row):
                        cells[i].text = str(val)

        # KM
        if st.session_state.get("ran_km") and "km_plot" in st.session_state:
            doc.add_heading("Kaplan-Meier Plot", level=1)
            tmp_km = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
            st.session_state["km_plot"].savefig(tmp_km.name)
            doc.add_picture(tmp_km.name, width=Inches(5))
            if "last_pvalue" in st.session_state:
                doc.add_paragraph(f"Log-Rank p-value: {st.session_state['last_pvalue']:.4f}")

        # Cox
        if st.session_state.get("ran_cox") and "cox_summary" in st.session_state:
            doc.add_heading("Cox Model Summary", level=1)
            df_cox = st.session_state["cox_summary"]
            table = doc.add_table(rows=1, cols=len(df_cox.columns))
            for i, col in enumerate(df_cox.columns):
                table.cell(0, i).text = col
            for row in df_cox.itertuples(index=False):
                cells = table.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = str(val)
            if "cox_plot" in st.session_state:
                doc.add_heading("Cox Plot", level=1)
                tmp_cox = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                st.session_state["cox_plot"].savefig(tmp_cox.name)
                doc.add_picture(tmp_cox.name, width=Inches(5))

        # Charts
        if st.session_state.get("charts"):
            for title, fig in st.session_state["charts"]:
                doc.add_heading(title, level=1)
                img_file = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                fig.savefig(img_file.name)
                doc.add_picture(img_file.name, width=Inches(5))

        # Save & download
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        with open(tmp.name, "rb") as f:
            st.download_button("📥 Download Final Word Report", f.read(), file_name=f"{file_name}_report.docx")

else:
    st.info("👈 Upload a dataset from the sidebar to begin.")
